"""
Testing iter_block_items()
"""
from __future__ import (
    absolute_import, division, print_function, unicode_literals
)

from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

#Variables & Data Structure

# From Paragraph
# note Overview and Notes fields here are used to store data after Contrast , Injection section in Word doc.
Protocol = {'Name': '', 'Overview': '', 'Notes': '', 'Extra': ''}
### Name
Name = {'Brand': '', 'Scanner': '', 'Number': ''}

# From Tables

"""Billing"""
Charges = ""

"""From Set up & Instructions"""
"""note: Topogram parameter is not present in Hierarchy table under Setup."""
Setup = {'Patient_Position': '',  'Breathing_Technique': '', 'Zero_location': '', 'Orientation': '',
         'Dose_Notification': '', 'Dose_Alert_Value': '', 'Instruction': '', 'Topogram': ''}

"""Contrast / Injection / Instructions"""
Medication = {'Oral_Contrast': '', 'IV_Contrast': '', 'Injection_rate': '', 'IV_Notes': '', 'Injection_rate_Notes': ''}

"""Note: Scout does not have Scout length parameter in hirearchy table."""
Scout = {'Scan_Type': '', 'kV': '', 'mA': '', 'Direction': '', 'Start': '', 'End': '', 'Scout_Plane': '', 'WW/WL': '',
         'Kernel': '', 'Destination': '', 'Scout_length': ''}

"""Following table is missing in Hirearchy Table: Routine Tab and Is picked up from Word Document."""
Routine = {'Eff_mAs': '', 'kV': '', 'Delay': '', 'Slice': '', 'Dose_Notification_Value': '', 'X-CARE': '',
           '4D_Range(Spital_Shuttle)': ''}

"""Note: Following parameters are just on and off parameters : CARE Dose4D, CARE kV, Dose_optimized_leve, Dual_energy,
   Hi_Res, Caradiac.
 Following parameters are missing in Hirerachy table: Number of scans, Feed, Ref QRM, Ref kV."""
Scan = {'Scan_Type': '', 'kV': '', 'mA': '', 'Rotation_time': '', 'Scan_coverage': '', 'Scan_Delay': '',
        'Scan_Direction': '', 'Thickness': '', 'Interval': '', 'Rotation_length': '','Detector_coverage': '',
        'Pitch': '', 'Speed': '', 'Tilt': '', 'SFOV': '', 'CARE Dose4D': '', 'CARE_kV': '', 'Dose_Optimized_level': '',
        'Dual_Energy': '', 'Hi_Res': '', 'Cardiac': '',
        'Description': '', 'Number_of_scans': '', 'Feed': '', 'Ref_QRM': '', 'Ref_kv': ''}

"""Note: Following parameters are just on and off parameters : SAFIRE, FAST."""
# Following parameters are missing in Hirerachy table: Slice, Window, FoV, Increment.
Recon = {'Description': '', 'DFOV': '', 'A/P_Center': '', 'R/L_Center': '', 'Thickness': '', 'Interval': '',
         'Algorithm': '', 'WW/WL': '', 'SAFIRE': '', 'SAFIRE_Strength': '', 'FAST': '','Recon_Kernel': '',
         'Recon_Slice_Data': '', 'Recon_Type': '', 'Recon_Region': '', 'Recon_Axis': '', '3D_Type': '',
         'Image_Order': '', 'Noise_Supression': '', 'Destinations': '',
         'Slice': '', 'Window': '', 'FoV': '', 'Increment': ''}

"""Following table is missing in Hirearchy Table: Recon tab yellow and Is picked up from Word Document."""
Bolus_Tracking = {'mAs': '', 'kV':'', 'Delay': '', 'Trigger(HU)': ''}

Series_Information = {'Series_Description': [], 'Slice': [], 'SAFIRE': [], 'SAFIRE_Strength': [], 'Algorithm': [],
                      'FAST': [], 'Window': [], 'FoV': [], 'Recon_job_type': [], 'Recon_Region': [],
                      'Recon_axis_for3D': [], 'Type_3D': [], 'Image_order': [], 'Increment': [], 'Destination': []}





def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
        # print(parent_elm.xml)
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

isFirstLine = True
isSetupInstructions = True
isRoutinekv = True
isRoutineSlice = True
isRoutineDelay = True
isScanDelay = True
isReconSafire = True
isReconStrength = True
isReconWindow =True
isReconFast = True
isBolusmas = True
isBoluskv = True
isBolusdelay = True



i = 1
document = Document('test.docx')
print (len(document.paragraphs))
print (len(document.tables))
for block in iter_block_items(document):
    #print (i)
    #i += 1
    if isinstance(block, Paragraph):
        #Code Goes here when there is Paragraph.
        if (isFirstLine):
            isFirstLine = False
            Protocol['Name'] =  block.text
            continue
        if 'siemens' in block.text.lower():
            Name['Brand'] = block.text.strip(' ()')
            continue
        if ':' in block.text.lower():
            Protocol['Extra'] = block.text
            continue
        print(block.text )
    else:
        #Code goes here when there is table.
        for row in block.rows:
            for cell in row.cells:
                # for paragraph in cell.paragraphs:
                if 'CHARGE' in cell.text.upper():
                    Charges = cell.text.split(':')[-1].strip()
                    continue
                if 'SET UP & INSTRUCTIONS' in cell.text and isSetupInstructions:
                    isSetupInstructions = False
                    Setup['Instruction'] = cell.text.strip().strip("SET UP & INSTRUCTIONS").strip().\
                        encode('ascii','replace')
                    continue
                if 'scout length' in cell.text.lower():
                    list = cell.text.split("\n")
                    scout_keywords = {'length': 'Scout_length', 'zero': 'Zero_location', 'type': 'Scan_Type'}
                    for i in list:
                        i =  i.encode('ascii','replace')
                        for word in scout_keywords:
                            if word in i.lower():
                                if word is 'zero':
                                    Setup[scout_keywords[word].encode('ascii','ignore')] = i.strip().split(":")[-1].\
                                        strip()
                                else:
                                    Scout[scout_keywords[word].encode('ascii','ignore')] = i.strip().split(":")[-1].\
                                        strip()
                        continue
                        # if 'length' in i.lower():
                        #     Scout['Scout_length'] = i.strip().split(":")[-1]
                        #     continue
                        # if 'zero' in i.lower():
                        #     Setup['Zero_location'] = i.strip().split(":")[-1]
                        #     continue
                        # if 'type' in i.lower():
                        #     Scout['Scan_Type'] = i.strip().split(":")[-1]
                        #     continue
                if 'patient position' in cell.text.lower():
                    list = cell.text.split("\n")
                    patient_keywords = {'patient': 'Patient_Position', 'breathing': 'Breathing_Technique',
                                        'topogram': 'Topogram'}
                    #print (list)
                    for i in list:
                        i =  i.encode('ascii','replace')
                        for word in scout_keywords:
                            if word in i.lower():
                                    Setup[scout_keywords[word].encode('ascii','ignore')] = i.strip().split(":")[-1].\
                                        strip()
                        continue
                        # if 'patient' in i.lower():
                        #     Setup['Patient_Position'] = i.strip().split(":")[-1]
                        #     continue
                        # if 'breathing' in i.lower():
                        #     Setup['Breathing_Technique'] = i.strip().split(":")[-1]
                        #     continue
                        # if 'topogram' in i.lower():
                        #     Setup['Topogram'] = i.strip().split(":")[-1]
                        #     continue
                if 'oral' in cell.text.lower():
                    Medication['Oral_Contrast'] = cell.text.split(":")[-1].strip("\n")
                    continue
                if 'iv:' in cell.text.lower():
                    list = cell.text.lower().split("special ins")
                    for i in list:
                        if 'iv:' in i:
                            Medication['IV_Contrast'] = i.split(":")[-1].strip()
                            continue
                        if 'ctions:' in i:
                            Medication['IV_Notes'] = i.split("ctions:")[-1].strip()
                            continue
                if 'injection' in cell.text.lower():
                    list = cell.text.lower().split("special ins")
                    for i in list:
                        if 'rate:' in i:
                            Medication['Injection_rate'] = i.split(":")[-1].strip()
                            continue
                        if 'ctions:' in i:
                            Medication['IV_Notes'] = i.split("ctions:")[-1].strip()
                            continue
                if 'protocol overview' in cell.text.lower():
                    Protocol['Overview'] = cell.text.lower().split("overview:")[-1].strip()
                    continue
                if 'special instruction:' in cell.text.lower():
                    Protocol['Notes'] = cell.text.lower().split("struction:")[-1].strip()
                    continue
                #Rotuine Tab
                routine_keywords = {
                    'eff. mas:': 'Eff_mAs', 'dose notification value': 'Dose_Notification_Value', 'x-care': 'X-CARE',
                    '4d range': '4D_Range(Spital_Shuttle)'
                }
                for word in routine_keywords:
                    if word in cell.text.lower():
                        Routine[routine_keywords[word]] = row.cells[1].text.strip()
                # if 'eff. mas:' in cell.text.lower():
                #     #print(len(row.cells))
                #     Routine['Eff_mAs'] = row.cells[1].text.strip()
                #     continue
                if 'kv:' in cell.text.lower() and isRoutinekv:
                    isRoutinekv = False
                    Routine['kV'] = row.cells[1].text.strip()
                    continue
                if 'delay:' in cell.text.lower() and isRoutineDelay:
                    isRoutineDelay = False
                    Routine['Delay'] = row.cells[1].text.strip()
                    continue
                if 'slice:' in cell.text.lower() and isRoutineSlice:
                    isRoutineSlice = False
                    Routine['Slice'] = row.cells[1].text.strip()
                    continue
                # if 'dose notification value' in cell.text.lower():
                #     Routine['Dose_Notification_Value'] = row.cells[1].text.strip()
                #     continue
                # if 'x-care' in cell.text.lower():
                #     Routine['X-CARE'] = row.cells[1].text.strip()
                #     continue
                # if '4d range' in cell.text.lower():
                #     Routine['4D_Range(Spital_Shuttle)'] = row.cells[1].text.strip()
                #     continue
                #Scan Tab.
                scan_keywords = {'care dose4d' : 'CARE Dose4D','number of scans': 'Number_of_scans',
                                 'care kv': 'CARE_kV', 'rotation time': 'Rotation_time', 'feed': 'Feed',
                                 'pitch' : 'Pitch', 'direction': 'Scan_Direction', 'ref qrm': 'Ref_QRM',
                                 'ref kv': 'Ref_kv', 'ref kv': 'Ref_kv', 'dose optimized level': 'Dose_Optimized_level'
                                 }
                for word in scan_keywords:
                    if word in cell.text.lower():
                        Scan[scan_keywords[word]] = row.cells[1].text.strip()
                # if 'care dose4d' in cell.text.lower():
                #     Scan['CARE Dose4D'] = row.cells[1].text.strip()
                #     continue
                # if 'number of scans' in cell.text.lower():
                #     Scan['Number_of_scans'] = row.cells[1].text.strip()
                #     continue
                # if 'care kv' in cell.text.lower():
                #     Scan['CARE_kV'] = row.cells[1].text.strip()
                #     continue
                #Rotation and Delay of Bolus are in same line needs to be taken care off
                # if 'rotation time' in cell.text.lower():
                #     Scan['Rotation_time'] = row.cells[1].text.strip()
                #     continue
                if 'delay' in cell.text.lower() and isBolusdelay:
                    isBolusdelay = False
                    Bolus_Tracking['Delay'] = row.cells[4].text.strip()
                    continue
                # if 'feed' in cell.text.lower():
                #     Scan['Feed'] = row.cells[1].text.strip()
                #     continue
                if 'DELAY' in cell.text.upper() and isScanDelay:
                    isScanDelay = False
                    Scan['Scan_Delay'] = row.cells[1].text.strip()
                    continue
                # if 'pitch' in cell.text.lower():
                #     Scan['Pitch'] = row.cells[1].text.strip()
                #     continue
                # if 'direction' in cell.text.lower():
                #     Scan['Scan_Direction'] = row.cells[1].text.strip()
                #     continue
                # if 'ref qrm' in cell.text.lower():
                #     Scan['Ref_QRM'] = row.cells[1].text.strip()
                #     continue
                # if 'ref kv' in cell.text.lower():
                #     Scan['Ref_kv'] = row.cells[1].text.strip()
                #     continue
                # if 'dose optimized level' in cell.text.lower():
                #     Scan['Dose_Optimized_level'] = row.cells[1].text.strip()
                #     continue
                #Recon Tab:
                if 'safire' in cell.text.lower() and isReconSafire:
                    isReconSafire = False
                    print(len(row.cells))
                    Recon['SAFIRE'] = row.cells[4].text.strip()
                    continue
                if 'strength (' in cell.text.lower() and isReconStrength:
                    isReconStrength = False
                    Recon['SAFIRE_Strength'] = row.cells[4].text.strip()
                    continue
                if 'fast' in cell.text.lower() and isReconFast:
                    isReconFast = False
                    Recon['FAST'] = row.cells[4].text.strip()
                    continue
                if 'window' in cell.text.lower() and isReconWindow:
                    isReconWindow = False
                    Recon['Window'] = row.cells[4].text.strip()
                    continue
                #Bolus Tracking
                # Can you isBolusmas Flag but is not required at moment since there are no conflicting words.
                bolus_keywords = {'mas:': 'mAs', 'kv:' : 'kv', 'trigger' : 'Trigger(HU)'}
                for word in bolus_keywords:
                    if word in cell.text.lower():
                        Bolus_Tracking[bolus_keywords[word]] = row.cells[4].text.strip()
                # if 'mas:' in cell.text.lower():
                #     Bolus_Tracking['mAs'] = row.cells[4].text.strip()
                #     continue
                # if 'kv:' in cell.text.lower():
                #     isBoluskv = False
                #     Bolus_Tracking['kv'] = row.cells[4].text.strip()
                #     continue
                # if 'trigger' in cell.text.lower():
                #     Bolus_Tracking['Trigger(HU)'] = row.cells[4].text.strip()
                #     continue
                # Series Information and Reformats
                SIR_keywords = {'series description': 'Series_Description', 'slice': 'Slice','safire': 'SAFIRE',
                                'safire strength': 'SAFIRE_Strength', 'algorithm': 'Algorithm','fast': 'FAST',
                                'window': 'Window', 'fov': 'FoV', 'recon job type': 'Recon_job_type',
                                'recon region': 'Recon_Region', 'recon axis:for 3d': 'Recon_axis_fro3D',
                                'type: 3d': 'Type_3D', 'image order': 'Image_order', 'increment': 'Increment',
                                'destination': 'Destination'}
                for word in SIR_keywords:
                    if word in cell.text.lower():

                        for i in range(1, len(row.cells)):
                            Series_Information[SIR_keywords[word].encode('ascii','ignore')].append(row.cells[i].text.strip())
                        continue
                # if 'series description' in cell.text.lower():
                #     for i in range(1, 4):
                #         Series_Information['Series_Description'].append(row.cells[i].text.strip())
                #     continue
                # if 'slice' in cell.text.lower():
                #     for i in range(1, 4):
                #         Series_Information['Slice'].append(row.cells[i].text.strip())
                #     continue
                # if 'safire' in cell.text.lower():
                #     for i in range(1, 4):
                #         Series_Information['SAFIRE'].append(row.cells[i].text.strip())
                #     continue
                # if 'safire strength' in cell.text.lower():
                #     for i in range(1, 4):
                #         Series_Information['SAFIRE_Strength'].append(row.cells[i].text.strip())
                #     continue
                # if 'algorithm' in cell.text.lower():
                #     for i in range(1, 4):
                #         Series_Information['Algorithm'].append(row.cells[i].text.strip())
                #     continue
                # if 'fast' in cell.text.lower():
                #     for i in range(1, 4):
                #         Series_Information['FAST'].append(row.cells[i].text.strip())
                #     continue
                # if 'window' in cell.text.lower():
                #     for i in range(1,4):
                #         Series_Information['Window'].append(row.cells[i].text.strip())
                #     continue
                # if 'fov' in cell.text.lower():
                #     for i in range(1, 4):
                #         Series_Information['FoV'].append(row.cells[i].text.strip())
                #     continue
                # if 'recon job type' in cell.text.lower():
                #     for i in range(1, 4):
                #         Series_Information['Recon_job_type'].append(row.cells[i].text.strip())
                #     continue
                # if 'recon region' in cell.text.lower():
                #     for i in range(1, 4):
                #         Series_Information['Recon_Region'].append(row.cells[i].text.strip())
                #     continue
                # if 'recon axis:for 3d' in cell.text.lower():
                #     for i in range(1, 4):
                #         Series_Information['Recon_axis_for3D'].append(row.cells[i].text.strip())
                #     continue
                # if 'type: 3d' in cell.text.lower():
                #     for i in range(1, 4):
                #         Series_Information['Type_3D'].append(row.cells[i].text.strip())
                #     continue
                # if 'image order' in cell.text.lower():
                #     for i in range(1, 4):
                #         Series_Information['Image_order'].append(row.cells[i].text.strip())
                #     continue
                # if 'increment' in cell.text.lower():
                #     for i in range(1, 4):
                #         Series_Information['Increment'].append(row.cells[i].text.strip())
                #     continue
                # if 'destination' in cell.text.lower():
                #     for i in range(1, 4):
                #         Series_Information['Destination'].append(row.cells[i].text.strip())
                #     continue

                print(cell.text)

print ("Protocol")
print (Protocol.items())
print ("Name")
print (Name.items())
print ("Charges")
print (Charges)
print ("Setup")
print (Setup.items())
print("Scout")
print (Scout.items())
print ("Medication")
print (Medication.items())
print ("Routine")
print (Routine.items())
print ("Scan")
print (Scan.items())
print ("Recon")
print (Recon.items())
print("Bolus Tracking")
print (Bolus_Tracking.items())
print ("Series Information")
print (Series_Information)