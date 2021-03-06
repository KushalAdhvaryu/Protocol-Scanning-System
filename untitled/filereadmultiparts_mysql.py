"""
Testing iter_block_items()
"""
from __future__ import (
    absolute_import, division, print_function, unicode_literals
)

from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import queryinsert_mysql as queryinsert
import mysql.connector
from mysql.connector import Error
import filereadmultiparts as data

#Variables & Data Structure

# From Paragraph
# note Overview and Notes fields here are used to store data after Contrast , Injection section in Word doc.
Protocol = {'Name': '', 'Overview': '', 'Notes': '', 'Extra': ''}
### Name
Name = {'Brand': '', 'Scanner': '', 'Number': ''}

# From Tables

"""Billing"""
Charges = {'charge': ''}

"""From Set up & Instructions"""
"""note: Topogram parameter is not present in Hierarchy table under Setup."""
Setup = {'Patient_Position': '',  'Breathing_Technique': '', 'Zero_location': '', 'Orientation': '',
         'Dose_Notification': '', 'Dose_Alert_Value': '', 'Instruction': '', 'Topogram': ''}

"""Contrast / Injection / Instructions"""
Medication = {'Oral_Contrast': '', 'IV_Contrast': '', 'Injection_rate': '', 'IV_Notes': '', 'Injection_rate_Notes': '',
              'Oral_Contrast_Note': ''}

"""Note: Scout does not have Scout length parameter in hirearchy table."""
Scout = {'Scan_Type': '', 'kV': '', 'mA': '', 'Direction': '', 'Start': '', 'End': '', 'Scout_Plane': '', 'WW/WL': '',
         'Kernel': '', 'Destination': '', 'Scout_length': ''}

"""Following table is missing in Hirearchy Table: Routine Tab and Is picked up from Word Document."""
Routine = {'Eff_mAs': [], 'kV': [], 'Delay': [], 'Slice': [], 'Dose_Notification_Value': [], 'X-CARE': [],
           '4D_Range(Spital_Shuttle)': []}

"""Note: Following parameters are just on and off parameters : CARE Dose4D, CARE kV, Dose_optimized_leve, Dual_energy,
   Hi_Res, Caradiac.
 Following parameters are missing in Hirerachy table: Number of scans, Feed, Ref QRM, Ref kV."""
Scan = {'Scan_Type': [], 'kV': [], 'mA': [], 'Rotation_time': [], 'Scan_coverage': [], 'Scan_Delay': [],
        'Scan_Direction': [], 'Thickness': [], 'Interval': [], 'Rotation_length': [],'Detector_coverage': [],
        'Pitch': [], 'Speed': [], 'Tilt': [], 'SFOV': [], 'CARE Dose4D': [], 'CARE_kV': [], 'Dose_Optimized_level': [],
        'Dual_Energy': [], 'Hi_Res': [], 'Cardiac': [],
        'Description': [], 'Number_of_scans': [], 'Feed': [], 'Ref_QRM': [], 'Ref_kv': []}

"""Note: Following parameters are just on and off parameters : SAFIRE, FAST."""
# Following parameters are missing in Hirerachy table: Slice, Window, FoV, Increment.
Recon = {'Description': [], 'DFOV': [], 'A/P_Center': [], 'R/L_Center': [], 'Thickness': [], 'Interval': [],
         'Algorithm': [], 'WW/WL': [], 'SAFIRE': [], 'SAFIRE_Strength': [], 'FAST': [],'Recon_Kernel': [],
         'Recon_Slice_Data': [], 'Recon_Type': [], 'Recon_Region': [], 'Recon_Axis': [], '3D_Type': [],
         'Image_Order': [], 'Noise_Supression': [], 'Destinations': [],
         'Slice': [], 'Window': [], 'FoV': [], 'Increment': [], 'Asir': []}

"""Following table is missing in Hirearchy Table: Recon tab yellow and Is picked up from Word Document."""
Bolus_Tracking = {'mAs': [], 'kV':[], 'Delay': [], 'Trigger(HU)': []}

# Series_Information = {'Series_Description': [], 'Slice': [], 'SAFIRE': [], 'SAFIRE_Strength': [], 'Algorithm': [],
#                       'FAST': [], 'Window': [], 'FoV': [], 'Recon_job_type': [], 'Recon_Region': [],
#                       'Recon_axis_for3D': [], 'Type_3D': [], 'Image_order': [], 'Increment': [], 'Destination': []}
# Recon : recon_type is same as Recon job type. recon_axis is same as Recon_Axis_for3D
Series_Information = {'Description': [], 'Slice': [], 'SAFIRE': [], 'SAFIRE_Strength': [], 'Algorithm': [],
                      'FAST': [], 'Window': [], 'FoV': [], 'Recon_Type': [], 'Recon_Region': [],
                      'Recon_axis': [], '3D_Type': [], 'Image_order': [], 'Increment': [], 'Destinations': []}

Extra_recon = []





def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
        # print(parent_elm.xml)
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)




def parse_document(inputfile):
    isFirstLine = True
    isSetupInstructions = True
    isRoutinekv = True
    isRoutineSlice = True
    isRoutineDelay = True
    isScanDelay = True
    isReconSafire = True
    isReconStrength = True
    isReconWindow = True
    isReconFast = True
    isBolusmas = True
    isBoluskv = True
    isBolusdelay = True
    ReconEffmasCount = 0
    i =1
    document = Document(inputfile)
    #print (len(document.paragraphs))
    #print (len(document.tables))
    for block in iter_block_items(document):
        #print (i)
        #i += 1
        if isinstance(block, Paragraph):
            #Code Goes here when there is Paragraph.
            if (isFirstLine):
                isFirstLine = False
                Protocol['Name'] =  block.text
                continue
            if 'siemens' in block.text.lower():
                Name['Brand'] = block.text.strip(' ()')
                continue
            if ':' in block.text.lower():
                Protocol['Extra'] = block.text
                continue
            #print(block.text )
        else:
            #Code goes here when there is table.
            for row in block.rows:
                for cell in row.cells:
                    # for paragraph in cell.paragraphs:
                    if 'CHARGE' in cell.text.upper():
                        Charges['charge'] = cell.text.split(':')[-1].strip()
                        continue
                    if 'SET UP & INSTRUCTIONS' in cell.text and isSetupInstructions:
                        isSetupInstructions = False
                        Setup['Instruction'] = cell.text.strip().strip("SET UP & INSTRUCTIONS").strip().\
                            encode('ascii','replace')
                        continue
                    if 'scout length' in cell.text.lower():
                        list = cell.text.split("\n")
                        scout_keywords = {'length': 'Scout_length', 'zero': 'Zero_location', 'type': 'Scan_Type'}
                        for i in list:
                            i =  i.encode('ascii','replace')
                            for word in scout_keywords:
                                if word in i.lower():
                                    if word is 'zero':
                                        Setup[scout_keywords[word].encode('ascii','ignore')] = i.strip().split(":")[-1].\
                                            strip()
                                    else:
                                        Scout[scout_keywords[word].encode('ascii','ignore')] = i.strip().split(":")[-1].\
                                            strip()
                            continue
                    if 'patient position' in cell.text.lower():
                        list = cell.text.split("\n")
                        patient_keywords = {'patient': 'Patient_Position', 'breathing': 'Breathing_Technique',
                                            'topogram': 'Topogram'}
                        #print (list)
                        for i in list:
                            i =  i.encode('ascii','replace')
                            for word in scout_keywords:
                                if word in i.lower():
                                        Setup[scout_keywords[word].encode('ascii','ignore')] = i.strip().split(":")[-1].\
                                            strip()
                            continue
                    if 'oral' in cell.text.lower():
                        list = cell.text.lower().split("special ins")
                        for i in list:
                            if 'rate:' in i:
                                Medication['Oral_Contrast'] = i.split(":")[-1].strip()
                                continue
                            if 'ctions:' in i:
                                Medication['Oral_Contrast_Note'] = i.split("ctions:")[-1].strip()
                                continue
                        # Medication['Oral_Contrast'] = cell.text.split(":")[-1].strip("\n")
                        # continue
                    if 'iv:' in cell.text.lower():
                        list = cell.text.lower().split("special ins")
                        for i in list:
                            if 'iv:' in i:
                                Medication['IV_Contrast'] = i.split(":")[-1].strip()
                                continue
                            if 'ctions:' in i:
                                Medication['IV_Notes'] = i.split("ctions:")[-1].strip()
                                continue
                    if 'injection' in cell.text.lower():
                        list = cell.text.lower().split("special ins")
                        for i in list:
                            if 'rate:' in i:
                                Medication['Injection_rate'] = i.split(":")[-1].strip()
                                continue
                            if 'ctions:' in i:
                                Medication['IV_Notes'] = i.split("ctions:")[-1].strip()
                                continue
                    if 'protocol overview' in cell.text.lower():
                        a,b = cell.text.lower().split("protocol overview:")
                        Protocol['Overview'] = a.strip() + " " + b.strip()
                        continue
                    if 'special instruction:' in cell.text.lower():
                        Protocol['Notes'] = cell.text.lower().split("struction:")[-1].strip()
                        continue
                    #Rotuine Tab
                    routine_keywords = {
                        'eff. mas:': 'Eff_mAs', 'dose notification value': 'Dose_Notification_Value', 'x-care': 'X-CARE',
                        '4d range': '4D_Range(Spital_Shuttle)'
                    }
                    for word in routine_keywords:
                        if word in cell.text.lower():
                            #This logic verifies that if document has multiple scan types like Chest and later abdoman it
                            # will have Eff. mAs: parameter twice , so for second occurance marking all flags to True to
                            #again read the words for setup.
                            if 'mas' in word:
                                ReconEffmasCount += 1
                                if ReconEffmasCount > 1:
                                    isRoutinekv = True
                                    isRoutineSlice = True
                                    isRoutineDelay = True
                                    isScanDelay = True
                                    isReconSafire = True
                                    isReconStrength = True
                                    isReconWindow = True
                                    isReconFast = True
                                    isBolusmas = True
                                    isBoluskv = True
                                    isBolusdelay = True
                            Routine[routine_keywords[word].encode('ascii','ignore')].append(row.cells[1].text.strip())
                    if 'kv:' in cell.text.lower() and isRoutinekv:
                        isRoutinekv = False
                        Routine['kV'].append(row.cells[1].text.strip())
                        continue
                    if 'delay:' in cell.text.lower() and isRoutineDelay:
                        isRoutineDelay = False
                        Routine['Delay'].append(row.cells[1].text.strip())
                        continue
                    if 'slice:' in cell.text.lower() and isRoutineSlice:
                        isRoutineSlice = False
                        Routine['Slice'].append(row.cells[1].text.strip())
                        continue
                    #Scan Tab.
                    scan_keywords = {'care dose4d' : 'CARE Dose4D','number of scans': 'Number_of_scans',
                                     'care kv': 'CARE_kV', 'rotation time': 'Rotation_time', 'feed': 'Feed',
                                     'pitch' : 'Pitch', 'direction': 'Scan_Direction', 'ref qrm': 'Ref_QRM',
                                     'ref kv': 'Ref_kv', 'ref kv': 'Ref_kv', 'dose optimized level': 'Dose_Optimized_level'
                                     }
                    for word in scan_keywords:
                        if word in cell.text.lower():
                            Scan[scan_keywords[word].encode('ascii','ignore')].append(row.cells[1].text.strip())
                    #Rotation and Delay of Bolus are in same line needs to be taken care off
                    if 'delay' in cell.text.lower() and isBolusdelay:
                        isBolusdelay = False
                        Bolus_Tracking['Delay'].append(row.cells[4].text.strip())
                        continue
                    if 'DELAY' in cell.text.upper() and isScanDelay:
                        isScanDelay = False
                        Scan['Scan_Delay'].append(row.cells[1].text.strip())
                        continue
                    #Recon Tab:
                    if 'safire' in cell.text.lower() and isReconSafire:
                        isReconSafire = False
                        #print(len(row.cells))
                        Recon['SAFIRE'].append(row.cells[4].text.strip())
                        continue
                    if 'strength (' in cell.text.lower() and isReconStrength:
                        isReconStrength = False
                        Recon['SAFIRE_Strength'].append(row.cells[4].text.strip())
                        continue
                    if 'fast' in cell.text.lower() and isReconFast:
                        isReconFast = False
                        Recon['FAST'].append(row.cells[4].text.strip())
                        continue
                    if 'window' in cell.text.lower() and isReconWindow:
                        isReconWindow = False
                        Recon['Window'].append(row.cells[4].text.strip())
                        continue
                    #Bolus Tracking
                    # Can you isBolusmas Flag but is not required at moment since there are no conflicting words.
                    bolus_keywords = {'mas:': 'mAs', 'kv:' : 'kV', 'trigger' : 'Trigger(HU)'}
                    for word in bolus_keywords:
                        if word in cell.text.lower():

                            if 'eff. mas:' in cell.text.lower():
                                #This corresponds to cell Eff. Mas due to similar matching so just ignore such words.
                                pass
                            else:
                                Bolus_Tracking[bolus_keywords[word].encode('ascii','ignore')].append(row.cells[4].text.strip())


                    if 'series information' in cell.text.lower():
                        isRoutinekv = False
                        isRoutineSlice = False
                        isRoutineDelay = False
                        isScanDelay = False
                        isReconSafire = False
                        isReconStrength = False
                        isReconWindow = False
                        isReconFast = False
                        isBolusmas = False
                        isBoluskv = False
                        isBolusdelay = False
                    # Series Information and Reformats
                    SIR_keywords = {'series description': 'Description', 'slice': 'Slice','safire:': 'SAFIRE',
                                    'safire strength': 'SAFIRE_Strength', 'algorithm': 'Algorithm','fast': 'FAST',
                                    'window': 'Window', 'fov': 'FoV', 'recon job type': 'Recon_Type',
                                    'recon region': 'Recon_Region', 'recon axis': 'Recon_axis',
                                    'type: 3d': '3D_Type', 'image order': 'Image_order', 'increment': 'Increment',
                                    'destination': 'Destinations'}
                    for word in SIR_keywords:
                        if word in cell.text.lower():

                            for i in range(1, len(row.cells)):
                                Series_Information[SIR_keywords[word].encode('ascii','ignore')].append(row.cells[i].text.strip())
                            continue
                if 'Recon ' in cell.text:
                    Extra_recon.append(cell.text.split(" ")[-1].strip())
                # print(cell.text)

def display_parseinfo():
    print ("----------------------------------Protocol-----------------------------------------------------------")
    for word in Protocol:
        print (word + " -> " + Protocol[word])
    print ("----------------------------------Name---------------------------------------------------------------")
    for word in Name:
        print (word + " -> " + Name[word])
    print ("----------------------------------Charges-----------------------------------------------------------")
    print (Charges['charge'])
    print ("----------------------------------Setup-----------------------------------------------------------")
    for word in Setup:
        print(word + " -> " + Setup[word])
    print("----------------------------------Scout-----------------------------------------------------------")
    for word in Scout:
        print(word + " -> " + Scout[word])
    print ("----------------------------------Medication-----------------------------------------------------------")
    for word in Medication:
        print(word + " -> " + Medication[word])
    print("----------------------------------Routine-----------------------------------------------------------")
    for word in Routine:
        print(word )
        print(Routine[word])
    print("----------------------------------Scan-----------------------------------------------------------")
    for word in Scan:
        print(word)
        print (Scan[word])
    print("----------------------------------Recon-----------------------------------------------------------")
    for word in Recon:
        print(word)
        print( Recon[word])
    print("----------------------------------Bolus Tracking-----------------------------------------------------------")
    for word in Bolus_Tracking:
        print(word)
        print( Bolus_Tracking[word])
    print("----------------------------------Series and information-----------------------------------------------------------")
    for word in Series_Information:
        print(word)
        print(Series_Information[word])
    print(
        "----------------------------------Extra Recon-----------------------------------------------------------")
    print (Extra_recon)

def getopenconnection():
    """ Connect to MySQL database """
    return mysql.connector.connect(host='127.0.0.1',
                                       database='protocol',
                                       user='root',
                                       password='1234')


if __name__ == "__main__":

    parse_document("test3.docx")

    # display_parseinfo()

    try:
        print("Getting connection to database")
        conn = getopenconnection()

        if conn.is_connected():
            print('Connected to MySQL database')

        #queryinsert.insert_protocol(Protocol,conn)
        #queryinsert.getProtocolid(conn)
        #queryinsert.insert_billing(conn, Charges)
        #queryinsert.insert_bolus_tracking(conn, Bolus_Tracking)
        #queryinsert.insert_medication(conn,Medication)
        # queryinsert.insert_name(conn, Name)
        # queryinsert.insert_recon_tab(conn, Recon)
        # queryinsert.insert_setup(conn, Setup)
        #queryinsert.insert_scout(conn, Scout)
        # queryinsert.insert_routine(conn, Routine)
        # queryinsert.insert_media(conn, "C:\\Users\\kadhvary\\PycharmProjects\\untitled\\output\\testpdf.pdf")
        #  queryinsert.read_media(conn ,"C:\\Users\\kadhvary\\PycharmProjects\\untitled\\pictures\\testpdf.pdf")
        # queryinsert.insert_scan(conn, Scan)
        queryinsert.insert_recon(conn, Series_Information)

        if conn:
            conn.close()

    except Exception as detail:
        print ("Something bad has happened!!! This is the error ==> ", detail)


