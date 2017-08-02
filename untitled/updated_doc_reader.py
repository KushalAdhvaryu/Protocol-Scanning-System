"""
Testing iter_block_items()
"""
from __future__ import (
    absolute_import, division, print_function, unicode_literals
)

from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import queryinsert
# import mysql.connector
import cx_Oracle
from mysql.connector import Error
import filereadmultiparts as data

#Variables & Data Structure

# From Paragraph
# note Overview and Notes fields here are used to store data after Contrast , Injection section in Word doc.
Protocol = {'Name': '', 'Overview': '', 'Notes': '', 'Extra': ''}
### Name
Name = {'Brand': '', 'Scanner': '', 'Number': ''}

# From Tables

"""Billing"""
Charges = {'charge': ''}

"""From Set up & Instructions"""
"""note: Topogram parameter is not present in Hierarchy table under Setup."""
Setup = {'Patient_Position': '',  'Breathing_Technique': '', 'Zero_location': '', 'Orientation': '',
         'Dose_Notification': '', 'Dose_Alert_Value': '', 'Instruction': '', 'Topogram': ''}

"""Contrast / Injection / Instructions"""
Medication = {'Oral_Contrast': '', 'IV_Contrast': '', 'Injection_rate': '', 'IV_Notes': '', 'Injection_rate_Notes': '',
              'Oral_Contrast_Note': ''}

"""Note: Scout does not have Scout length parameter in hirearchy table."""
Scout = {'Scan_Type': '', 'kV': '', 'mA': '', 'Direction': '', 'Start': '', 'End': '', 'Scout_Plane': '', 'WW/WL': '',
         'Kernel': '', 'Destination': '', 'Scout_length': ''}

"""Following table is missing in Hirearchy Table: Routine Tab and Is picked up from Word Document."""
Routine = {'Eff_mAs': [], 'kV': [], 'Delay': [], 'Slice': [], 'Dose_Notification_Value': [], 'X-CARE': [],
           '4D_Range(Spital_Shuttle)': []}

"""Note: Following parameters are just on and off parameters : CARE Dose4D, CARE kV, Dose_optimized_leve, Dual_energy,
   Hi_Res, Caradiac.
 Following parameters are missing in Hirerachy table: Number of scans, Feed, Ref QRM, Ref kV."""
Scan = {'Scan_Type': [], 'kV': [], 'mA': [], 'Rotation_time': [], 'Scan_coverage': [], 'Scan_Delay': [],
        'Scan_Direction': [], 'Thickness': [], 'Interval': [], 'Rotation_length': [],'Detector_coverage': [],
        'Pitch': [], 'Speed': [], 'Tilt': [], 'SFOV': [], 'CARE Dose4D': [], 'CARE_kV': [], 'Dose_Optimized_level': [],
        'Dual_Energy': [], 'Hi_Res': [], 'Cardiac': [],
        'Description': [], 'Number_of_scans': [], 'Feed': [], 'Ref_QRM': [], 'Ref_kv': []}

"""Note: Following parameters are just on and off parameters : SAFIRE, FAST."""
# Following parameters are missing in Hirerachy table: Slice, Window, FoV, Increment.
Recon = {'Description': [], 'DFOV': [], 'A/P_Center': [], 'R/L_Center': [], 'Thickness': [], 'Interval': [],
         'Algorithm': [], 'WW/WL': [], 'SAFIRE': [], 'SAFIRE_Strength': [], 'FAST': [],'Recon_Kernel': [],
         'Recon_Slice_Data': [], 'Recon_Type': [], 'Recon_Region': [], 'Recon_Axis': [], '3D_Type': [],
         'Image_Order': [], 'Noise_Supression': [], 'Destinations': [],
         'Slice': [], 'Window': [], 'FoV': [], 'Increment': [], 'Asir': []}

"""Following table is missing in Hirearchy Table: Recon tab yellow and Is picked up from Word Document."""
Bolus_Tracking = {'mAs': [], 'kV':[], 'Delay': [], 'Trigger(HU)': []}

# Series_Information = {'Series_Description': [], 'Slice': [], 'SAFIRE': [], 'SAFIRE_Strength': [], 'Algorithm': [],
#                       'FAST': [], 'Window': [], 'FoV': [], 'Recon_job_type': [], 'Recon_Region': [],
#                       'Recon_axis_for3D': [], 'Type_3D': [], 'Image_order': [], 'Increment': [], 'Destination': []}
# Recon : recon_type is same as Recon job type. recon_axis is same as Recon_Axis_for3D
Series_Information = {'Description': [], 'Slice': [], 'SAFIRE': [], 'SAFIRE_Strength': [], 'Algorithm': [],
                      'FAST': [], 'Window': [], 'FoV': [], 'Recon_Type': [], 'Recon_Region': [],
                      'Recon_axis': [], '3D_Type': [], 'Image_order': [], 'Increment': [], 'Destinations': []}

Extra_recon = []






def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
        # print(parent_elm.xml)
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)




def parse_document(inputfile):
    isFirstLine = True
    isSetupInstructions = True


    isRoutinetab = True
    isSeriesInfo = False



    ReconEffmasCount = 0
    i =1
    document = Document(inputfile)
    #print (len(document.paragraphs))
    #print (len(document.tables))
    for block in iter_block_items(document):
        #print (i)
        #i += 1
        if isinstance(block, Paragraph):
            #Code Goes here when there is Paragraph.
            if (isFirstLine):
                isFirstLine = False
                Protocol['Name'] =  block.text
                continue
            if 'siemens' in block.text.lower():
                Name['Brand'] = block.text.strip(' ()')
                continue
            if ':' in block.text.lower():
                Protocol['Extra'] = block.text
                continue
            #print(block.text )
        else:
            #Code goes here when there is table.
            for row in block.rows:
                for cell in row.cells:
                    # for paragraph in cell.paragraphs:
                    if 'CHARGE' in cell.text.upper():
                        Charges['charge'] = cell.text.split(':')[-1].strip()
                        continue
                    if 'SET UP & INSTRUCTIONS' in cell.text and isSetupInstructions:
                        isSetupInstructions = False
                        setupcolParse(block)
                    # if 'SET UP & INSTRUCTIONS' in cell.text and isSetupInstructions:
                    #     isSetupInstructions = False
                    #     Setup['Instruction'] = cell.text.strip().strip("SET UP & INSTRUCTIONS").strip().\
                    #         encode('ascii','replace')
                    #     continue
                    # if 'scout length' in cell.text.lower():
                    #     list = cell.text.split("\n")
                    #     scout_keywords = {'length': 'Scout_length', 'zero': 'Zero_location', 'type': 'Scan_Type'}
                    #     for i in list:
                    #         i =  i.encode('ascii','replace')
                    #         for word in scout_keywords:
                    #             if word in i.lower():
                    #                 if word is 'zero':
                    #                     Setup[scout_keywords[word].encode('ascii','ignore')] = i.strip().split(":")[-1].\
                    #                         strip()
                    #                 else:
                    #                     Scout[scout_keywords[word].encode('ascii','ignore')] = i.strip().split(":")[-1].\
                    #                         strip()
                    #         continue
                    # if 'patient position' in cell.text.lower():
                    #     list = cell.text.split("\n")
                    #     patient_keywords = {'patient': 'Patient_Position', 'breathing': 'Breathing_Technique',
                    #                         'topogram': 'Topogram'}
                    #     #print (list)
                    #     for i in list:
                    #         i =  i.encode('ascii','replace')
                    #         for word in patient_keywords:
                    #             if word in i.lower():
                    #                     Setup[patient_keywords[word].encode('ascii','ignore')] = i.strip().split(":")[-1].\
                    #                         strip()
                    #         continue
                    # if 'oral' in cell.text.lower():
                    #     list = cell.text.lower().split("special ins")
                    #     for i in list:
                    #         if 'rate:' in i:
                    #             Medication['Oral_Contrast'] = i.split(":")[-1].strip()
                    #             continue
                    #         if 'ctions:' in i:
                    #             Medication['Oral_Contrast_Note'] = i.split("ctions:")[-1].strip()
                    #             continue
                    #     # Medication['Oral_Contrast'] = cell.text.split(":")[-1].strip("\n")
                    #     # continue
                    # if 'iv:' in cell.text.lower():
                    #     list = cell.text.lower().split("special ins")
                    #     for i in list:
                    #         if 'iv:' in i:
                    #             Medication['IV_Contrast'] = i.split(":")[-1].strip()
                    #             continue
                    #         if 'ctions:' in i:
                    #             Medication['IV_Notes'] = i.split("ctions:")[-1].strip()
                    #             continue
                    # if 'injection rate' in cell.text.lower():
                    #     list = cell.text.lower().split("special ins")
                    #     for i in list:
                    #         if 'rate:' in i:
                    #             Medication['Injection_rate'] = i.split(":")[-1].strip()
                    #             continue
                    #         if 'ctions:' in i:
                    #             Medication['Injection_rate_Notes'] = i.split("ctions:")[-1].strip()
                    #             continue
                    # if 'protocol overview' in cell.text.lower() or \
                    #                 'protocol overview and instructions' in cell.text.lower():
                    #     a,b = cell.text.lower().split("protocol overview")
                    #     Protocol['Overview'] = a.strip() + " " + b.strip()
                    #     continue
                    # if 'special instruction:' in cell.text.lower():
                    #     Protocol['Notes'] = cell.text.lower().split("struction:")[-1].strip()
                    #     continue
                    #Rotuine tab
                    if 'routine tab' in cell.text.lower():
                        isSeriesInfo = False
                        if isRoutinetab:
                            tabcolParse(block)
                            # tabsParsing(block, ReconEffmasCount)
                            isRoutinetab = False


                    if 'series information' in cell.text.lower():
                        isSeriesInfo = True
                        isRoutinetab = True

                    if isSeriesInfo:
                    # Series Information and Reformats

                        SIR_keywords = {'series description': 'Description', 'slice': 'Slice','safire:': 'SAFIRE',
                                    'safire strength': 'SAFIRE_Strength', 'algorithm': 'Algorithm','fast': 'FAST',
                                    'window': 'Window', 'fov': 'FoV', 'recon job type': 'Recon_Type',
                                    'recon region': 'Recon_Region', 'recon axis': 'Recon_axis',
                                    'type: 3d': '3D_Type', 'image order': 'Image_order', 'increment': 'Increment',
                                    'destination': 'Destinations'}
                        for word in SIR_keywords:
                            if word in cell.text.lower():

                                for i in range(1, len(row.cells)):
                                    Series_Information[SIR_keywords[word].encode('ascii','ignore')].append(row.cells[i].text.strip())
                                continue


                if 'Recon ' in cell.text:
                    # print("#$%^$#%^#&&^%&$%&%$&^%&^^")
                    # print(cell.text)
                    # print("#$%^$#%^#&&^%&$%&%$&^%&^^")
                    Extra_recon.append(cell.text.split(" ")[-1].strip())
                # print(cell.text)


def setupcolParse(block):
    additionalinfo = []
    for row in block.rows:
        for cell in row.cells:
            if 'SET UP & INSTRUCTIONS' in cell.text:
                if cell.text.strip().strip("SET UP & INSTRUCTIONS").strip().encode('ascii', 'replace') not in \
                        Setup['Instruction']:
                    Setup['Instruction'] += cell.text.strip().strip("SET UP & INSTRUCTIONS").strip(). \
                    encode('ascii', 'replace') + " "
                    continue
                continue
            if 'scout length' in cell.text.lower():
                list = cell.text.split("\n")
                scout_keywords = {'length': 'Scout_length', 'zero': 'Zero_location', 'type': 'Scan_Type'}
                for i in list:
                    i = i.encode('ascii', 'replace')
                    for word in scout_keywords:
                        if word in i.lower():
                            if word is 'zero':
                                Setup[scout_keywords[word].encode('ascii', 'ignore')] = i.strip().split(":")[-1]. \
                                    strip()
                            else:
                                Scout[scout_keywords[word].encode('ascii', 'ignore')] = i.strip().split(":")[-1]. \
                                    strip()
                continue
            if 'patient position' in cell.text.lower():
                list = cell.text.split("\n")
                patient_keywords = {'patient': 'Patient_Position', 'breathing': 'Breathing_Technique',
                                    'topogram': 'Topogram'}
                # print (list)
                for i in list:
                    i = i.encode('ascii', 'replace')
                    for word in patient_keywords:
                        if word in i.lower():
                            Setup[patient_keywords[word].encode('ascii', 'ignore')] = i.strip().split(":")[-1]. \
                                strip()
                continue
            if 'oral' in cell.text.lower():
                list = cell.text.lower().split("special ins")
                for i in list:
                    if 'rate:' in i:
                        Medication['Oral_Contrast'] = i.split(":")[-1].strip()
                        continue
                    if 'ctions:' in i:
                        Medication['Oral_Contrast_Note'] = i.split("ctions:")[-1].strip()
                        continue
                        # Medication['Oral_Contrast'] = cell.text.split(":")[-1].strip("\n")
                continue
            if 'iv:' in cell.text.lower():
                list = cell.text.lower().split("special ins")
                for i in list:
                    if 'iv:' in i:
                        Medication['IV_Contrast'] = i.split(":")[-1].strip()
                        continue
                    if 'ctions:' in i:
                        Medication['IV_Notes'] = i.split("ctions:")[-1].strip()
                        continue
                continue
            if 'injection rate' in cell.text.lower():
                list = cell.text.lower().split("special ins")
                for i in list:
                    if 'rate:' in i:
                        Medication['Injection_rate'] = i.split(":")[-1].strip()
                        continue
                    if 'ctions:' in i:
                        Medication['Injection_rate_Notes'] = i.split("ctions:")[-1].strip()
                        continue
                continue
            if 'protocol overview' in cell.text.lower() or \
                            'protocol overview and instructions' in cell.text.lower():
                # a, b = cell.text.lower().split("protocol overview")
                # Protocol['Overview'] = a.strip() + " " + b.strip()
                Protocol['Overview'] = cell.text.lower()
                continue
            if 'special instruction:' in cell.text.lower():
                if cell.text.lower().strip() not in Setup['Instruction']:
                    Setup['Instruction'] += cell.text.lower().strip()
                    continue
                continue
            if cell.text not in additionalinfo:
                if 'CONTRAST' not in cell.text and '/' not in cell.text and 'INJECTION' not in cell.text:
                    additionalinfo.append(cell.text)
                    print(additionalinfo)
    for element in additionalinfo:
            Setup['Instruction'] += element + " "

def tabcolParse(block):
    # print("Here......................................................................................")
    routineTabctr = 0
    reconTabctr = 0
    Lbls = []
    Vals = []
    RoutLbls = []
    RoutVals = []

    for col in block.columns:
        for cell in col.cells:
            # print (cell.text)
            if 'routine tab' in cell.text.lower():
                routineTabctr +=1
            if routineTabctr == 1:
                Lbls.append(cell.text.lower().strip())
            if routineTabctr == 2:
                Vals.append(cell.text.lower().strip())
            if 'recon tab' in cell.text.lower():
                reconTabctr +=1
            if reconTabctr == 1:
                RoutLbls.append(cell.text.lower().strip())
            if reconTabctr == 2:
                RoutVals.append(cell.text.lower().strip())
    bolus_tracking_idx = -1
    # print (Lbls.index('routine tab:'))
    routine_tab_idx = Lbls.index('routine tab:')
    # print(Lbls.index('scan tab:'))
    scan_tab_idx = Lbls.index('scan tab:')
    # print(RoutLbls.index('recon tab:'))
    recon_tab_idx = RoutLbls.index('recon tab:')
    if 'bolus tracking' in RoutLbls:
        # print(RoutLbls.index('bolus tracking'))
        bolus_tracking_idx = RoutLbls.index('bolus tracking')
    # for i in range(len(Lbls)):
    #     if Lbls[i] != '':
    #         print (str(i) +" " + Lbls[i] + " -> " + Vals[i])
    # for i in range(len(Lbls)):
    #     if RoutLbls[i] != '':
    #         print (str(i)+ " " + RoutLbls[i] + " -> " + RoutVals[i])

    routine_keywords = {
        'eff. mas:': 'Eff_mAs','kv:':'kV' ,'dose notification value': 'Dose_Notification_Value', 'x-care': 'X-CARE',
        '4d range': '4D_Range(Spital_Shuttle)','delay:':'Delay','slice:':'Slice'
    }
    for word in routine_keywords:
        for i in range(routine_tab_idx+1,scan_tab_idx):
            if word in Lbls[i]:
                Routine[routine_keywords[word].encode('ascii', 'ignore')].append(Vals[i])

    scan_keywords = {'care dose4d': 'CARE Dose4D', 'number of scans': 'Number_of_scans',
                     'care kv': 'CARE_kV', 'rotation time': 'Rotation_time', 'feed': 'Feed',
                     'pitch': 'Pitch', 'direction': 'Scan_Direction', 'ref qrm': 'Ref_QRM',
                     'ref kv': 'Ref_kv', 'dose optimized level': 'Dose_Optimized_level','delay':'Scan_Delay'
                     }
    for word in scan_keywords:
        for i in range(scan_tab_idx+1,len(Lbls)):
            if word in Lbls[i]:
                Scan[scan_keywords[word].encode('ascii', 'ignore')].append(Vals[i])

    recon_keywords = {'safire:':'SAFIRE','strength':'SAFIRE_Strength','fast':'FAST', 'window':'Window'
                    }
    if bolus_tracking_idx == -1:
        end_idx = len(RoutLbls)
    else:
        end_idx = bolus_tracking_idx

    for word in recon_keywords:
        for i in range(recon_tab_idx+1,end_idx):
            if word in RoutLbls[i]:
                Recon[recon_keywords[word].encode('ascii', 'ignore')].append(RoutVals[i])

    bolus_keywords = {'mas:': 'mAs', 'kv:': 'kV', 'trigger': 'Trigger(HU)','delay':'Delay'}
    if bolus_tracking_idx != -1:
        for word in bolus_keywords:
            for i in range(bolus_tracking_idx + 1, len(RoutLbls)):
                if word in RoutLbls[i]:
                    Bolus_Tracking[bolus_keywords[word].encode('ascii', 'ignore')].append(RoutVals[i])
    # print("Here......................................................................................")



def display_parseinfo():
    print ("----------------------------------Protocol-----------------------------------------------------------")
    for word in Protocol:
        print (word + " -> " + Protocol[word])
    print ("----------------------------------Name---------------------------------------------------------------")
    for word in Name:
        print (word + " -> " + Name[word])
    print ("----------------------------------Charges-----------------------------------------------------------")
    print (Charges['charge'])
    print ("----------------------------------Setup-----------------------------------------------------------")
    for word in Setup:
        print(word + " -> " + Setup[word])
    print("----------------------------------Scout-----------------------------------------------------------")
    for word in Scout:
        print(word + " -> " + Scout[word])
    print ("----------------------------------Medication-----------------------------------------------------------")
    for word in Medication:
        print(word + " -> " + Medication[word])
    print("----------------------------------Routine-----------------------------------------------------------")
    for word in Routine:
        print(word )
        print(Routine[word])
    print("----------------------------------Scan-----------------------------------------------------------")
    for word in Scan:
        print(word)
        print (Scan[word])
    print("----------------------------------Recon-----------------------------------------------------------")
    for word in Recon:
        print(word)
        print( Recon[word])
    print("----------------------------------Bolus Tracking-----------------------------------------------------------")
    for word in Bolus_Tracking:
        print(word)
        print( Bolus_Tracking[word])
    print("----------------------------------Series and information-----------------------------------------------------------")
    for word in Series_Information:
        print(word)
        print(Series_Information[word])
    print(
        "----------------------------------Extra Recon-----------------------------------------------------------")
    print (Extra_recon)

def getopenconnection():
    """ Connect to MySQL database 
    return mysql.connector.connect(host='127.0.0.1',
                                       database='protocol',
                                       user='root',
                                       password='1234')
    """
    return cx_Oracle.connect("ANONYMOUS", "mayo", "mayo")

if __name__ == "__main__":

    parse_document("test5.docx")

    display_parseinfo()

    try:
        print("Getting connection to database")
        conn = getopenconnection()
        print ('Connected to Oracle database Version '+conn.version)

        source = 'word'
        queryinsert.insert_protocol(Protocol,conn,source)
        queryinsert.getProtocolid(conn)
        queryinsert.insert_billing(conn, Charges)
        queryinsert.insert_bolus_tracking(conn, Bolus_Tracking)
        queryinsert.insert_medication(conn,Medication)
        queryinsert.insert_name(conn, Name)
        queryinsert.insert_recon_tab(conn, Recon)
        queryinsert.insert_setup(conn, Setup,Protocol)
        queryinsert.insert_scout(conn, Scout)
        queryinsert.insert_routine(conn, Routine)
        queryinsert.insert_media(conn, "C:\\Users\\kadhvary\\PycharmProjects\\untitled\\pictures\\testpdf.pdf", Protocol)
        queryinsert.insert_scan(conn, Scan)
        queryinsert.insert_recon(conn, Series_Information)

        if conn:
            conn.close()

    except Exception as detail:
        print ("Something bad has happened!!! This is the error ==> ", detail)


