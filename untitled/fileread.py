"""
Testing iter_block_items()
"""
from __future__ import (
    absolute_import, division, print_function, unicode_literals
)

from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

#Variables & Data Structure

# From Paragraph
# note Overview and Notes fields here are used to store data after Contrast , Injection section in Word doc.
Protocol = {'Name': '', 'Overview': '', 'Notes': '', 'Extra': ''}
### Name
Name = {'Brand': '', 'Scanner': '', 'Number': ''}

# From Tables

"""Billing"""
Charges = ""

"""From Set up & Instructions"""
"""note: Topogram parameter is not present in Hierarchy table under Setup."""
Setup = {'Patient_Position': '',  'Breathing_Technique': '', 'Zero_location': '', 'Orientation': '',
         'Dose_Notification': '', 'Dose_Alert_Value': '', 'Instruction': '', 'Topogram': ''}

"""Contrast / Injection / Instructions"""
Medication = {'Oral_Contrast': '', 'IV_Contrast': '', 'Injection_rate': '', 'IV_Notes': '', 'Injection_rate_Notes': ''}

"""Note: Scout does not have Scout length parameter in hirearchy table."""
Scout = {'Scan_Type': '', 'kV': '', 'mA': '', 'Direction': '', 'Start': '', 'End': '', 'Scout_Plane': '', 'WW/WL': '',
         'Kernel': '', 'Destination': '', 'Scout_length': ''}

"""Following table is missing in Hirearchy Table: Routine Tab and Is picked up from Word Document."""
Routine = {'Eff_mAs': '', 'kV': '', 'Delay': '', 'Slice': '', 'Dose_Notification_Value': '', 'X-CARE': '',
           '4D_Range(Spital_Shuttle)': ''}

"""Note: Following parameters are just on and off parameters : CARE Dose4D, CARE kV, Dose_optimized_leve, Dual_energy,
   Hi_Res, Caradiac.
 Following parameters are missing in Hirerachy table: Number of scans, Feed, Ref QRM, Ref kV."""
Scan = {'Scan_Type': '', 'kV': '', 'mA': '', 'Rotation_time': '', 'Scan_coverage': '', 'Scan_Delay': '',
        'Scan_Direction': '', 'Thickness': '', 'Interval': '', 'Rotation_length': '','Detector_coverage': '',
        'Pitch': '', 'Speed': '', 'Tilt': '', 'SFOV': '', 'CARE Dose4D': '', 'CARE_kV': '', 'Dose_Optimized_level': '',
        'Dual_Energy': '', 'Hi_Res': '', 'Cardiac': '',
        'Description': '', 'Number_of_scans': '', 'Feed': '', 'Ref_QRM': '', 'Ref_kv': ''}

"""Note: Following parameters are just on and off parameters : SAFIRE, FAST."""
# Following parameters are missing in Hirerachy table: Slice, Window, FoV, Increment.
Recon = {'Description': '', 'DFOV': '', 'A/P_Center': '', 'R/L_Center': '', 'Thickness': '', 'Interval': '',
         'Algorithm': '', 'WW/WL': '', 'SAFIRE': '', 'SAFIRE_Strength': '', 'FAST': '','Recon_Kernel': '',
         'Recon_Slice_Data': '', 'Recon_Type': '', 'Recon_Region': '', 'Recon_Axis': '', '3D_Type': '',
         'Image_Order': '', 'Noise_Supression': '', 'Destinations': '',
         'Slice': '', 'Window': '', 'FoV': '', 'Increment': ''}

"""Following table is missing in Hirearchy Table: Recon tab yellow and Is picked up from Word Document."""
Bolus_Tracking = {'mAs': '', 'kV':'', 'Delay': '', 'Trigger(HU)': ''}

Series_Information = {'Series_Description': [], 'Slice': [], 'SAFIRE': [], 'SAFIRE_Strength': [], 'Algorithm': [],
                      'FAST': [], 'Window': [], 'FoV': [], 'Recon_job_type': [], 'Recon_Region': [],
                      'Recon_axis_for3D': [], 'Type_3D': [], 'Image_order': [], 'Increment': [], 'Destination': []}





def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
        # print(parent_elm.xml)
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

isFirstLine = True
isSetupInstructions = True
isRoutinekv = True
isRoutineSlice = True
isRoutineDelay = True
isScanDelay = True
isReconSafire = True
isReconStrength = True
isReconWindow =True
isReconFast = True
isBolusmas = True
isBoluskv = True
isBolusdelay = True



i = 1
document = Document('test1.docx')
#print (len(document.paragraphs))
#print (len(document.tables))
for block in iter_block_items(document):
    #print (i)
    #i += 1
    if isinstance(block, Paragraph):
        #Code Goes here when there is Paragraph.
        if (isFirstLine):
            isFirstLine = False
            Protocol['Name'] =  block.text
            continue
        if 'siemens' in block.text.lower():
            Name['Brand'] = block.text.strip(' ()')
            continue
        if ':' in block.text.lower():
            Protocol['Extra'] = block.text
            continue
        #print(block.text )
    else:
        #Code goes here when there is table.
        for row in block.rows:
            for cell in row.cells:
                # for paragraph in cell.paragraphs:
                if 'CHARGE' in cell.text.upper():
                    Charges = cell.text.split(':')[-1].strip()
                    continue
                if 'SET UP & INSTRUCTIONS' in cell.text and isSetupInstructions:
                    isSetupInstructions = False
                    Setup['Instruction'] = cell.text.strip().strip("SET UP & INSTRUCTIONS").strip().\
                        encode('ascii','replace')
                    continue
                if 'scout length' in cell.text.lower():
                    list = cell.text.split("\n")
                    scout_keywords = {'length': 'Scout_length', 'zero': 'Zero_location', 'type': 'Scan_Type'}
                    for i in list:
                        i =  i.encode('ascii','replace')
                        for word in scout_keywords:
                            if word in i.lower():
                                if word is 'zero':
                                    Setup[scout_keywords[word].encode('ascii','ignore')] = i.strip().split(":")[-1].\
                                        strip()
                                else:
                                    Scout[scout_keywords[word].encode('ascii','ignore')] = i.strip().split(":")[-1].\
                                        strip()
                        continue
                if 'patient position' in cell.text.lower():
                    list = cell.text.split("\n")
                    patient_keywords = {'patient': 'Patient_Position', 'breathing': 'Breathing_Technique',
                                        'topogram': 'Topogram'}
                    #print (list)
                    for i in list:
                        i =  i.encode('ascii','replace')
                        for word in scout_keywords:
                            if word in i.lower():
                                    Setup[scout_keywords[word].encode('ascii','ignore')] = i.strip().split(":")[-1].\
                                        strip()
                        continue
                if 'oral' in cell.text.lower():
                    Medication['Oral_Contrast'] = cell.text.split(":")[-1].strip("\n")
                    continue
                if 'iv:' in cell.text.lower():
                    list = cell.text.lower().split("special ins")
                    for i in list:
                        if 'iv:' in i:
                            Medication['IV_Contrast'] = i.split(":")[-1].strip()
                            continue
                        if 'ctions:' in i:
                            Medication['IV_Notes'] = i.split("ctions:")[-1].strip()
                            continue
                if 'injection' in cell.text.lower():
                    list = cell.text.lower().split("special ins")
                    for i in list:
                        if 'rate:' in i:
                            Medication['Injection_rate'] = i.split(":")[-1].strip()
                            continue
                        if 'ctions:' in i:
                            Medication['IV_Notes'] = i.split("ctions:")[-1].strip()
                            continue
                if 'protocol overview' in cell.text.lower():
                    Protocol['Overview'] = cell.text.lower().split("overview:")[-1].strip()
                    continue
                if 'special instruction:' in cell.text.lower():
                    Protocol['Notes'] = cell.text.lower().split("struction:")[-1].strip()
                    continue
                #Rotuine Tab
                routine_keywords = {
                    'eff. mas:': 'Eff_mAs', 'dose notification value': 'Dose_Notification_Value', 'x-care': 'X-CARE',
                    '4d range': '4D_Range(Spital_Shuttle)'
                }
                for word in routine_keywords:
                    if word in cell.text.lower():
                        Routine[routine_keywords[word]] = row.cells[1].text.strip()
                if 'kv:' in cell.text.lower() and isRoutinekv:
                    isRoutinekv = False
                    Routine['kV'] = row.cells[1].text.strip()
                    continue
                if 'delay:' in cell.text.lower() and isRoutineDelay:
                    isRoutineDelay = False
                    Routine['Delay'] = row.cells[1].text.strip()
                    continue
                if 'slice:' in cell.text.lower() and isRoutineSlice:
                    isRoutineSlice = False
                    Routine['Slice'] = row.cells[1].text.strip()
                    continue
                #Scan Tab.
                scan_keywords = {'care dose4d' : 'CARE Dose4D','number of scans': 'Number_of_scans',
                                 'care kv': 'CARE_kV', 'rotation time': 'Rotation_time', 'feed': 'Feed',
                                 'pitch' : 'Pitch', 'direction': 'Scan_Direction', 'ref qrm': 'Ref_QRM',
                                 'ref kv': 'Ref_kv', 'ref kv': 'Ref_kv', 'dose optimized level': 'Dose_Optimized_level'
                                 }
                for word in scan_keywords:
                    if word in cell.text.lower():
                        Scan[scan_keywords[word]] = row.cells[1].text.strip()
                #Rotation and Delay of Bolus are in same line needs to be taken care off
                if 'delay' in cell.text.lower() and isBolusdelay:
                    isBolusdelay = False
                    Bolus_Tracking['Delay'] = row.cells[4].text.strip()
                    continue
                if 'DELAY' in cell.text.upper() and isScanDelay:
                    isScanDelay = False
                    Scan['Scan_Delay'] = row.cells[1].text.strip()
                    continue
                #Recon Tab:
                if 'safire' in cell.text.lower() and isReconSafire:
                    isReconSafire = False
                    #print(len(row.cells))
                    Recon['SAFIRE'] = row.cells[4].text.strip()
                    continue
                if 'strength (' in cell.text.lower() and isReconStrength:
                    isReconStrength = False
                    Recon['SAFIRE_Strength'] = row.cells[4].text.strip()
                    continue
                if 'fast' in cell.text.lower() and isReconFast:
                    isReconFast = False
                    Recon['FAST'] = row.cells[4].text.strip()
                    continue
                if 'window' in cell.text.lower() and isReconWindow:
                    isReconWindow = False
                    Recon['Window'] = row.cells[4].text.strip()
                    continue
                #Bolus Tracking
                # Can you isBolusmas Flag but is not required at moment since there are no conflicting words.
                bolus_keywords = {'mas:': 'mAs', 'kv:' : 'kv', 'trigger' : 'Trigger(HU)'}
                for word in bolus_keywords:
                    if word in cell.text.lower():
                        Bolus_Tracking[bolus_keywords[word]] = row.cells[4].text.strip()
                # Series Information and Reformats
                SIR_keywords = {'series description': 'Series_Description', 'slice': 'Slice','safire': 'SAFIRE',
                                'safire strength': 'SAFIRE_Strength', 'algorithm': 'Algorithm','fast': 'FAST',
                                'window': 'Window', 'fov': 'FoV', 'recon job type': 'Recon_job_type',
                                'recon region': 'Recon_Region', 'recon axis': 'Recon_axis_for3D',
                                'type: 3d': 'Type_3D', 'image order': 'Image_order', 'increment': 'Increment',
                                'destination': 'Destination'}
                for word in SIR_keywords:
                    if word in cell.text.lower():

                        for i in range(1, len(row.cells)):
                            Series_Information[SIR_keywords[word].encode('ascii','ignore')].append(row.cells[i].text.strip())
                        continue
                #print(cell.text)

print ("----------------------------------Protocol-----------------------------------------------------------")
for word in Protocol:
    print (word + " -> " + Protocol[word])
print ("----------------------------------Name---------------------------------------------------------------")
for word in Name:
    print (word + " -> " + Name[word])
print ("----------------------------------Charges-----------------------------------------------------------")
print (Charges)
print ("----------------------------------Setup-----------------------------------------------------------")
for word in Setup:
    print(word + " -> " + Setup[word])
print("----------------------------------Scout-----------------------------------------------------------")
for word in Scout:
    print(word + " -> " + Scout[word])
print ("----------------------------------Medication-----------------------------------------------------------")
for word in Medication:
    print(word + " -> " + Medication[word])
print("----------------------------------Routine-----------------------------------------------------------")
for word in Routine:
    print(word + " -> " + Routine[word])
print("----------------------------------Scan-----------------------------------------------------------")
for word in Scan:
    print(word + " -> " + Scan[word])
print("----------------------------------Recon-----------------------------------------------------------")
for word in Recon:
    print(word + " -> " + Recon[word])
print("----------------------------------Bolus Tracking-----------------------------------------------------------")
for word in Bolus_Tracking:
    print(word + " -> " + Bolus_Tracking[word])
print("----------------------------------Series and information-----------------------------------------------------------")
for word in Series_Information:
    print(word)
    print(Series_Information[word])


